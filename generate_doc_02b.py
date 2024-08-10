from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()


def add_content_ia_b(document):
# Título
    document.add_heading('Infraestrutura de Nuvem do Azure para o Aplicativo Siteblitz', 0)

    # Seção 1: Armazenamento de Objetos
    document.add_heading('Armazenamento de Objetos', 1)
    document.add_paragraph('O armazenamento de objetos foi criado para fornecer uma solução de armazenamento escalável e econômica para dados não estruturados, como imagens, vídeos e logs. Ele foi configurado para dimensionar e manipular automaticamente grandes quantidades de dados.')
    document.add_paragraph('')

    # Seção 2: Backend do Siteblitz
    document.add_heading('Backend do Siteblitz', 1)
    document.add_paragraph('O backend do Siteblitz foi criado usando o Azure Kubernetes Service (AKS) e o k3s. Ele foi configurado para dimensionar e manipular automaticamente alto tráfego.')
    document.add_paragraph('')

    # Seção 3: Máquinas Virtuais
    document.add_heading('Máquinas Virtuais', 1)
    document.add_paragraph('As máquinas virtuais foram criadas para hospedar o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 4: Dados
    document.add_heading('Dados', 1)
    document.add_paragraph('O armazenamento de dados foi configurado usando bancos de dados SQL do Azure.')
    document.add_paragraph('')

    # Seção 5: Kubernetes (k3s)
    document.add_heading('Kubernetes (k3s)', 1)
    document.add_paragraph('O Kubernetes foi usado para gerenciar e orquestrar o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 6: Pods
    document.add_heading('Pods', 1)
    document.add_paragraph('Os pods foram criados para executar o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 7: Serviço
    document.add_heading('Serviço', 1)
    document.add_paragraph('Os serviços foram criados para expor o aplicativo Siteblitz e outros serviços à Internet.')
    document.add_paragraph('')

    # Seção 8: Ingress
    document.add_heading('Ingress', 1)
    document.add_paragraph('O Ingress foi criado para rotear o tráfego para o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 9: Rede
    document.add_heading('Rede', 1)
    document.add_paragraph('Uma rede virtual foi criada para isolar o aplicativo Siteblitz e outros serviços da Internet.')
    document.add_paragraph('')

    # Seção 10: Zona DNS
    document.add_heading('Zona DNS', 1)
    document.add_paragraph('Uma zona DNS foi criada para fornecer um nome de domínio personalizado para o aplicativo Siteblitz.')
    document.add_paragraph('')

    # Seção 11: Gateway de Rede Virtual
    document.add_heading('Gateway de Rede Virtual', 1)
    document.add_paragraph('Um gateway de rede virtual foi criado para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 12: Gateway de Rede Local
    document.add_heading('Gateway de Rede Local', 1)
    document.add_paragraph('Um gateway de rede local foi criado para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 13: Balanceamento de Carga (Gateway de Aplicativo)
    document.add_heading('Balanceamento de Carga (Gateway de Aplicativo)', 1)
    document.add_paragraph('Um gateway de aplicativo foi criado para fornecer balanceamento de carga e roteamento de tráfego para o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 14: Segurança
    document.add_heading('Segurança', 1)
    document.add_paragraph('A segurança foi configurada usando o centro de segurança do Azure para fornecer proteção contra ameaças, avaliação de vulnerabilidades e aplicação de políticas de segurança.')
    document.add_paragraph('')

    # Seção 15: WAF
    document.add_heading('WAF', 1)
    document.add_paragraph('Um firewall de aplicativo da Web (WAF) foi criado para fornecer proteção contra ataques de aplicativo da Web.')
    document.add_paragraph('')

    # Seção 16: Registro de Contêiner
    document.add_heading('Registro de Contêiner', 1)
    document.add_paragraph('Um registro de contêiner foi criado para armazenar e gerenciar imagens de contêiner para o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 17: Gateway VPN
    document.add_heading('Gateway VPN', 1)
    document.add_paragraph('Um gateway VPN foi criado para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 18: VPN Site-to-site (IPsec)
    document.add_heading('VPN Site-to-site (IPsec)', 1)
    document.add_paragraph('Uma VPN site-to-site foi criada para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 19: VPN Site-to-site (IPsec) ETRO CONSTRUCTION com METATRON
    document.add_heading('VPN Site-to-site (IPsec) ETRO CONSTRUCTION com METATRON', 1)
    document.add_paragraph('Uma VPN site-to-site foi criada para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Seção 20: VPN Site-to-site (IPsec) ETRO CONSTRUCTION com AUDAZ TECNOLOGIA
    document.add_heading('VPN Site-to-site (IPsec) ETRO CONSTRUCTION com AUDAZ TECNOLOGIA', 1)
    document.add_paragraph('Uma VPN site-to-site foi criada para fornecer conectividade segura entre o aplicativo Siteblitz e outros serviços.')
    document.add_paragraph('')

    # Salvar o documento