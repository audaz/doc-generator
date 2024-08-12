import time
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


import re
import os, inspect
from win32com import client
import win32com
from pathlib import Path

from docx.oxml import OxmlElement as OE
from docx.oxml.ns import qn

def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000
    # return (section.page_width - section.left_margin - section.right_margin) 


def add_list_of_table(run):
    fldChar = OE('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true')
    instrText = OE('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Table"'  #"Table" of list of table and "Figure" for list of figure
    fldChar2 = OE('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OE('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)
    
    fldChar4 = OE('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)

# def create_header(document):

#     section = document.sections[0]
#     header = section.header
#     paragraph = header.paragraphs[0]
#     paragraph.text = "Left Text\tCenter Text\tRight Text"
#     paragraph.style = document.styles["Header"]


def add_content(document):
    # Add the title
    document.add_heading('As-built Document of XCorp\'s Cloud Infrastructure', 0)

    # Add the subtitle
    document.add_paragraph('This document provides an overview of XCorp\'s cloud infrastructure, which includes Kubernetes, KEDA, NGINX Ingress Controller, AWS CloudFront, AWS S3, AWS ECR, AWS Router53, Oracle Cloud Instance, RabbitMQ, Redis for caching, MySQL Heatwave, Postgres Database, ElasticSearch for logging and APM, and PRTG.' )

    # Add a new section
    document.add_heading('1. Architecture', level=1)

    # Add a new paragraph
    document.add_paragraph('XCorp\'s cloud infrastructure is designed to be highly available, scalable, and secure. The infrastructure is built on top of AWS and Oracle Cloud, with Kubernetes serving as the container orchestration platform.')

    # Add a new section
    document.add_heading('1.1 Kubernetes', level=2)
    document.add_paragraph('Kubernetes is used to manage and orchestrate containers across multiple nodes in the cluster. The cluster consists of multiple worker nodes, each running multiple pods, which are the smallest deployable units of computing in Kubernetes.')

    # Add a new section
    document.add_heading('1.2 KEDA', level=2)
    document.add_paragraph('KEDA (Kubernetes Event-driven Autoscaling) is used to automatically scale pods based on the number of events or messages received by the pod. This allows for fine-grained control over the scaling of applications and ensures that resources are used efficiently.')

    # Add a new section
    document.add_heading('1.3 NGINX Ingress Controller', level=2)
    document.add_paragraph('NGINX Ingress Controller is used to route traffic to the appropriate pods in the cluster. It acts as a reverse proxy and load balancer, ensuring that traffic is distributed evenly across the cluster.')

    # Add a new section
    document.add_heading('1.4 AWS CloudFront', level=2)
    document.add_paragraph('AWS CloudFront is used as a content delivery network (CDN) to cache static content and serve it to users from the nearest edge location. This helps to improve the performance of the application and reduce the load on the backend servers.')

    # Add a new section
    document.add_heading('1.5 AWS S3', level=2)
    document.add_paragraph('AWS S3 is used as a storage service for static content, such as images and videos. It provides high availability and durability, and can be used to store large amounts of data at a low cost.')


    # Add a new section
    document.add_heading('1.6 AWS ECR', level=2)
    document.add_paragraph('AWS ECR (Elastic Container Registry) is used as a container registry to store and manage container images. It provides secure and scalable storage for container images, and can be used to deploy containers to the cluster.')

    # Add a new section
    document.add_heading('1.7 AWS Router53', level=3)
    document.add_paragraph('AWS Router53 is used as a DNS service to route traffic to the appropriate IP address. It provides high availability and scalability, and can be used to manage DNS records for the application.')

    # Add a new section
    document.add_heading('2 Oracle Cloud Instance', level=1)
    document.add_paragraph('Oracle Cloud Instance is used as a compute service')




def add_arquitetura(document):
    with open("arquitetura.md", "r", encoding="utf-8") as file:
        # user_message += file.read().replace("\n", "")
        content = file.read()

    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.startswith('<<'):
            img = paragraph.replace('<<','').replace('>>','')
            print(f"{img}")
            p1 = document.add_picture(f'assets\\{img}', width=Inches(7))
        elif paragraph.startswith('1.') or paragraph.startswith('2.') or paragraph.startswith('3.') or paragraph.startswith('4.') or paragraph.startswith('5.') or paragraph.startswith('6.') or paragraph.startswith('7.') or paragraph.startswith('8.') or paragraph.startswith('9.') or paragraph.startswith('10.') or paragraph.startswith('11.') or paragraph.startswith('12.') or paragraph.startswith('13.') or paragraph.startswith('14.') or paragraph.startswith('15.') or paragraph.startswith('16.') or paragraph.startswith('17.') or paragraph.startswith('18.') or paragraph.startswith('19.') or paragraph.startswith('20.') or paragraph.startswith('21.') or paragraph.startswith('22.') or paragraph.startswith('23.'):
            document.add_heading(paragraph, level=1)
        else:
            p1 = document.add_paragraph(paragraph)
            logo_run = p1.add_run()
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   


class REMatcher(object):
    def __init__(self, matchstring):
        self.matchstring = matchstring

    def match(self,regexp):
        self.rematch = re.match(regexp, self.matchstring)
        return bool(self.rematch)

    def group(self,i):
        return self.rematch.group(i)


def test_marko():
    pass
    # mark = marko.parse(content)
    # list_children = mark.children
    # for child in list_children:
    #     print(child)
    #     time.sleep(1)
    # print(mark.__dict__)
    # time.sleep(100)
