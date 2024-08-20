# -*- coding: utf-8 -*-
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


import re
import os, inspect
from win32com import client
import win32com
from pathlib import Path

from docx.oxml import OxmlElement as OE
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from src.gen_pi_html2img import gen_code_image
from src.gen_pillow_mac_style import gen_code_image_windows
from src.gen_pygments import generate_pygments

from render_marko import parse_content_from_markdown



dbg = 0
# base_dir = 'D:\\Dev\\doc-generator\\'
# base_dir = f"{os.path.dirname(os.path.realpath(__file__))}\\"

base_dir = os.path.abspath(os.path.realpath(__file__))
base_dir, file_script = os.path.split(base_dir)

def update_toc(docx_file):
    print(f"funcion update_toc - dispatch\n")
    word = win32com.client.DispatchEx("Word.Application")
    print(f"open        {docx_file}\n")
    doc = word.Documents.Open(f"{docx_file}")
    toc_count = doc.TablesOfContents.Count
    print(f'toc count={toc_count}')
    if int(toc_count) == 1:
        print(f"update\n")
        doc.TablesOfContents(1).Update()
    elif int(toc_count) == 1:
        print(f"Create\n")
        doc.TablesOfContents().Create()
    else:
        print('Error: não existe apenas 1 indice')
    print(f"close\n")
    doc.Close(SaveChanges=True)
    print(f"quit\n")
    word.Quit()


def create_front_page(empresa, title, directory, document):
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')

    document.add_picture(f'{base_dir}\\assets\\imagem_capa_{empresa}.png', width=Mm(210))

    p = document.add_paragraph('')
    p = document.add_paragraph('')

    paragraph_00 = document.add_paragraph()
    logo_run = paragraph_00.add_run(title)
    logo_run.font.bold = True
    logo_run.font.size = Pt(16)
    paragraph_00.alignment = WD_ALIGN_PARAGRAPH.CENTER    

    p = document.add_paragraph('')
    p = document.add_paragraph('')

    p2 = document.add_paragraph('')
    logo_run_02 = p2.add_run()
    l_file = Path(f"{base_dir}{directory}\\logo.png")
    if l_file.is_file():
        logo_run_02.add_picture(f"{base_dir}{directory}\\logo.png", width=Inches(1.5)) 
    else:
        p = document.add_paragraph('')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    # p  = document.add_paragraph('')

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    if empresa == 'nikos':
        logo_run_03 = p3.add_run('www.nikos.com.br')
    else:
        logo_run_03 = p3.add_run('www.audaztecnologia.com.br')
    logo_run_03.font.bold = True
    logo_run_03.font.size = Pt(14)

    document.add_page_break()


def create_blank_toc(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph_01 = doc.add_paragraph('SUMÁRIO')
    title_style = paragraph_01.style
    title_style.font.name = "Arial"
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_style.font.bold = True
    title_style.font.size = Pt(11)
    title_style.element.xml
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial")    

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

    doc.add_page_break()



def format_title_and_subtitles(document):
    # Set the font for the paragraphs
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)
    document.styles['Normal'].font.bold = False
    document.styles['Normal'].font.italic = False
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal'].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set the font for the titles and subtitles
    document.styles['Title'].font.name = 'Arial'
    document.styles['Title'].font.size = Pt(20)
    document.styles['Title'].font.bold = True
    document.styles['Title'].font.italic = False
    document.styles['Title'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Title'].element.xml
    rFonts = document.styles['Title'].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial")

    document.styles['Heading 1'].font.name = 'Arial'
    document.styles['Heading 1'].font.size = Pt(12)
    document.styles['Heading 1'].font.bold = True
    document.styles['Heading 1'].font.italic = False
    document.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Heading 1'].element.xml
    rFonts = document.styles['Heading 1'].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial")    

    document.styles['Heading 2'].font.name = 'Arial'
    document.styles['Heading 2'].font.size = Pt(12)
    document.styles['Heading 2'].font.bold = True
    document.styles['Heading 2'].font.italic = False
    document.styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Heading 2'].element.xml
    rFonts = document.styles['Heading 2'].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial")     

    document.styles['Heading 3'].font.name = 'Arial'
    document.styles['Heading 3'].font.size = Pt(12)
    document.styles['Heading 3'].font.bold = True
    document.styles['Heading 3'].font.italic = False
    document.styles['Heading 3'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Heading 3'].element.xml
    rFonts = document.styles['Heading 3'].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial") 

    document.styles['Heading 4'].font.name = 'Arial'
    document.styles['Heading 4'].font.size = Pt(12)
    document.styles['Heading 4'].font.bold = True
    document.styles['Heading 4'].font.italic = False
    document.styles['Heading 4'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Heading 4'].element.xml
    rFonts = document.styles['Heading 4'].element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Arial")     





def add_content_direct_from_markdown(document,directory, file_path_directory, file):
    # with open(file, "r" , encoding="utf-8") as file:
    if dbg>=0: print(f"\nFILE: {file}")
    with open(f"{file}", "r" , encoding="utf-8") as file:
    # with open(f"{base_dir}{file}", "r" , encoding="utf-8") as file:
        # user_message += file.read().replace("\n", "")
        content = file.read()
    # if dbg>=0: print(content)
    # exit(0)
    paragraphs = content.split('\n')
    inside_table = False
    table_created = False
    table_columns_total = 0
    table_items = ()
    table = [None] * 2000
    table_count = 0
    inside_code = False
    code_creating = False
    code_text = [None] * 100
    code_id = 0
    paragraph_id = 0

    for paragraph in paragraphs:
        paragraph_id = paragraph_id + 1
        # if dbg>0 :print(type(paragraph))
        if dbg>=1: print(f"\n{paragraph}")

        # if paragraph.startswith('!['):
        if m := re.search(r'^```(.*)',      paragraph):    
            print('CODE DETECTED...')
            print(f"{m.group(1)}")
            inside_code = True            
            if code_creating == True:
                code_creating = False            
                inside_code = False            
                max, lines, img_file = create_image_code(code_text[code_id], directory,paragraph_id)
                print(img_file)
                time.sleep(5)
                size_max = max * 1.7
                if size_max >= 85:
                    size_max = 150 
                document.add_picture(img_file, width=Mm( size_max ))
            else:
                code_creating = True            
                code_id = code_id + 1
                code_text[code_id] = ""
            inside_table = False
            # time.sleep(1)
        elif inside_code:
            print(f'inside_code code_id={code_id}')
            code_text[code_id] += f"{paragraph}\n"
            # time.sleep(0.1)
        elif m := re.search(r'^\s*(!\[.*)$', paragraph): # Parse images   
            img = m.group(1)
            img = paragraph.replace('[','').replace(']','').replace(r'!','').replace(r'\s+','').replace(' ','').replace(' ','').replace(' ','')
            img = img.split('(')[0]
            print(f"IMG: {img}")
            pi1 = document.add_paragraph('')
            pi1.alignment = WD_ALIGN_PARAGRAPH.CENTER            
            run_pi1 = pi1.add_run()
            # run_pi1.add_picture(f'{base_dir}{directory}\\{img}' )
            run_pi1.add_picture(f'{file_path_directory}\\{img}' )
            inside_table = False
            inside_code = False            
        # https://stackoverflow.com/questions/2554185/match-groups-in-python
        elif m := re.search(r'^\s*([0-9]+\.[0-9]+\.[0-9]+.*)$', paragraph):    
            print(f"REGEX_03b_{m.group(0)}")
            print(f"REGEX_03b_{m.group(1)}")
            # print(f"REGEX_03b_{m.group(2)}")
            document.add_heading(f"{m.group(1)}", level=3)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^([0-9]+\.[0-9]+\.[0-9]+.*)', paragraph):    
            document.add_heading(f"{m.group(1)}", level=3)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^\s*([0-9]+\.[0-9]+.*)',    paragraph):    
            document.add_heading(f"{m.group(1)}", level=2)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^\s*([0-9]+\.\s+.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=1)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^---\s*$',      paragraph):    
            print('add page break')
            document.add_page_break()
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^######\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=6)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^#####\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=5)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^####\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=4)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^###\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=3)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^##\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=2)
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^#\s(.*)',      paragraph):    
            document.add_heading(f"{m.group(1)}", level=1)
            inside_table = False
            inside_code = False            
        # parse * isso é uma lista
        elif m := re.search(r'^\*(.*)',      paragraph):    
            document.add_paragraph(m.group(1), style='List Bullet') 
            inside_table = False
            inside_code = False            
        elif m := re.search(r'^<!--',      paragraph):    
            inside_table = False
            inside_code = False            
            next
        elif m := re.search(r'^\|\s?---+(.*)',      paragraph):    
            if dbg >=1: print('TABLE HEADER DETECTED...')
            if dbg >=1: print(f"{m.group(1)}")
            inside_table = True
            inside_code = False            
            if table_created:
                if dbg >=1: print('Tabela já existe...')
            else:
                table_count = table_count + 1
                table[table_count] = document.add_table(rows=1, cols=len(re.findall(r'\|', m.group(1))))
                # table = document.add_table(rows=0, cols=10)
                table[table_count].style = 'Table Grid'
                table[table_count].allow_autofit = True
                table_created = True
        elif m := re.search(r'^\|(.*)',      paragraph):    
            if dbg >=1: print('TABLE DETECTED...')
            if dbg >=1: print(f"{m.group(1)}")
            inside_table = True
            inside_code = False            
            table_columns_total = len(re.findall(r'\|', m.group(1)))
            if table_created:
                if dbg >=1: print('Tabela já existe...')
            else:
                table_count = table_count + 1
                table[table_count] = document.add_table(rows=0, cols=table_columns_total)
                # table = document.add_table(rows=0, cols=10)
                table[table_count].style = 'Table Grid'
                table[table_count].allow_autofit = True
                table_created = True
            if dbg >=1: print(f"table_columns_total:{table_columns_total}")
            if dbg >=1: print(f"col_count = {len(table[table_count].columns)}")
            row_itens = m.group(1).split('|')
            cells = table[table_count].add_row().cells


            for i in range(table_columns_total):
                if dbg>=1: print(f"i:{i}  row_itens[i]:{row_itens[i]}")
                if c := re.search(r'(.*)\[size=(\d+)\]', row_itens[i]):
                    text = c.group(1)
                    text = parse_hyperlink(text)
                    size = int(c.group(2))
                    print(f'Tamanho customoziado de tabela encontrado. size={size}')
                    for cell in table[table_count].columns[i].cells:
                        cell.width = Cm(size)
                    cells[i].text = text
                    # time.sleep(1)
                else:
                    text = row_itens[i] 
                    text = parse_hyperlink(text)
                    text = parse_hyperlink(text)
                    text = parse_hyperlink(text)
                    cells[i].text = text 
                    # Apply color only on row[0]
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
                    table[table_count].rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
                    table[table_count].rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255) 
        else:
            text = parse_hyperlink(paragraph)
            p1 = document.add_paragraph(text)
            logo_run = p1.add_run()
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY               
            p1.paragraph_format.space_after = Pt(2)
            inside_table = False
            inside_code = False            

        if inside_table == False:
            table_created = False
def parse_hyperlink(text):
    # if t := re.search(r'(.*)\[().*)\]\(.*)(\)(.*)', text):
    if t := re.search(r'\[\[([^\]]+)\]\(#([^\)]+)\)\]\(#\[[^\]]+\]\(#([^\)]+)\)\)' , text):
        if dbg>=1: print(t.group(0))
        if dbg>=1: print(t.group(1))
        if dbg>=1: print(t.group(2))
        if dbg>=1: print(t.group(3))
        new_text = f"{t.group(1)}"  
        link = t.group(3)
        if dbg>=1: print(f'HYPERLIK DETECTADO NA REGEX1 {new_text} link:{link}')
        if dbg>=1: time.sleep(2)
        return new_text        
    elif t := re.search(r'^(.*)\[(.*)\]\((.*)\)(.*)', text):
        if dbg>=1: print(t.group(0))
        if dbg>=1: print(t.group(1))
        if dbg>=1: print(t.group(2))
        if dbg>=1: print(t.group(3))
        if dbg>=1: print(t.group(4))
        new_text = f"{t.group(1)}{t.group(2)}{t.group(4)}"  
        link = t.group(3)
        if dbg>=1: print(f'HYPERLIK DETECTADO {new_text} link:{link}')
        # if dbg>=1: time.sleep(2)
        return new_text
    else:
        return text


def create_image_code(code_text,directory,img_id):
    (max, lines, html_file) = generate_pygments(code_text,directory,img_id)
    # time.sleep(10)
    img_file = gen_code_image(max, lines, html_file)
    # gen_code_image_windows()
    return max, lines, img_file

def generate_content(content_file, directory, empresa='audaz', title='' , custom_file='',  convert_to_pdf=True, create_cover=True, create_toc=True):
    # script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    # file_path = os.path.join(script_dir, f"{directory}\\{file_name}")
    if custom_file != '' :
        print('Custom file detectado')
        file_with_output_path = f"{custom_file[:-3]}.docx"
        file_with_output_path = os.path.abspath(file_with_output_path)
        file_path_directory , file_path_filename = os.path.split(file_with_output_path)
        file_path_directory = f"{file_path_directory}" + "\\"

        file_with_output_path = f"{file_path_directory}\\output\\{file_path_filename}"
        # file_name = custom_file 
        file_name = file_with_output_path 
        directory = file_path_directory
        # content_file = custom_file
        content_file = file_with_output_path
    else:
        file_name = f"{content_file[:-3]}.docx"
        file_path_directory = directory
        file_with_output_path = os.path.join(base_dir, f"{directory}\\output\\{file_name}")
        file_path_directory_new , file_path_filename = os.path.split(file_with_output_path)

    print(f"Opening file directory {directory}...")
    print(f"Opening file file_path_directory {file_path_directory}...")
    print(f"Opening file file_with_output_path {file_with_output_path}...")
    print(f"Opening file file_path docx (old)  {file_with_output_path}...")
    print(f"generating file file_name .md {file_name}...")
    print(f"content_file .md {content_file}...")
    time.sleep(1)

    if not os.path.exists(f"{file_path_directory}\\output"):
        os.makedirs(f"{file_path_directory}\\output")

    # Create a new document
    document = Document()

    import docx 
    styles_element = document.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(docx.oxml.shared.qn('w:val'),'pt-BR')

    if title == '':
        title = file_path_filename[:-5]
        print(f'Redefinindo título:  {title}  ...')
        # time.sleep(20)

    if create_cover :
        create_front_page(empresa, title, directory, document)

    section = document.sections[0]
    # Page A4 - https://stackoverflow.com/questions/43724030/how-to-change-page-size-to-a4-in-python-docx
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(4)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)    

    header = section.header
    pH = header.paragraphs[0]
    logo_run = pH.add_run()
    logo_run.add_picture(f"{base_dir}\\assets\\logo_{empresa}.png", width=Inches(1.5))    
    pH.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text_run = pH.add_run()
    text_run.text = '\t' +  "__________________________________________________________________________________\n"
    text_run.style.font.color.rgb = RGBColor(58, 19, 19)

    text_rum_2 = pH.add_run()
    text_run_2 = '\t'

    if create_toc:
        create_blank_toc(document)

    format_title_and_subtitles(document)

    if custom_file != '':
        add_content_direct_from_markdown(document,directory,file_path_directory, custom_file)
    else:
        add_content_direct_from_markdown(document,directory,file_path_directory, f"{directory}\\{content_file}")
    # parse_content_from_markdown(document,directory,f"{directory}\\{content_file}")


    # print(f"file_path_directory: {file_path_directory}")
    # print(f"file_path: {file_path}")
    # time.sleep(50)

    document.save(file_with_output_path)

    if create_toc:
        update_toc(file_with_output_path)

    if convert_to_pdf:
        print("convert to pdf")
        pdf_file = f"{file_with_output_path[:-5]}.pdf"
        print(f"pdf_file = {pdf_file}")
        convert(file_with_output_path, f"{pdf_file}")

def main(args):
    print(f"args2:{args}")
    print(f"args2.file:{args.file}")    
    file = args.file
    empresa = args.empresa
    # file = file.replace('\\','\\\\')
    file = file.replace('\"','')
    print(f"args3:{file}")
    if file:
        # exit(0)
        # if os.path.exists(f"\"{file}\""):
        if os.path.exists(file):
            # exit(0)
            if empresa:
                print(type(file))
                print(type(empresa))
                print(file)
                print(empresa)
                generate_content("nao_se_aplica.md", directory='data\\nao_sem_aplica',  convert_to_pdf=True, custom_file=file, empresa=args.empresa)
            else:
                generate_content("nao_se_aplica.md", directory='data\\nao_se_aplica',  convert_to_pdf=True, custom_file=file)

        else:
            print(f"arquivo não encontrado. {args.file}")
            exit(1)
    else:
        generate_content("manual.md", title='Manual para criação de novos ambientes (DevOps)',  convert_to_pdf=False)
        # generate_content("proposta_vxlan.md", directory='data\\audaz_modelo_proposta',  convert_to_pdf=True, create_cover=False)
        # generate_content("arquitetura.md", directory='data\\metatron_arquitetura',  convert_to_pdf=True)
        # generate_content("as-built-02-ptbr.md",title='Infraestrutura em nuvem (DevOps)',  directory='data\\metatron_infraestrutura',  convert_to_pdf=False)
        # generate_content("esteira_ci_cd.md", directory='data\\metratron_cid',  convert_to_pdf=True)
        # generate_content("infra.md", title='Infra TEN Meetings', directory='data\\ten_meetings_infra',  convert_to_pdf=True)
        # generate_content("Relatório - FinOps TEN.md", title='Relatório - FinOps TEN', directory='data\\ten_relatorio_finops',  convert_to_pdf=True, create_cover=True, create_toc=False)
        # generate_content("dr.md", empresa='nikos', title='Plano de recuperação de Desastre', directory='data\\nikos_dr',  convert_to_pdf=True)
        # generate_content("okit_markdown_completo.md", empresa='nikos', title='Documentação da infraestrutura em nuvem', directory='data\\nikos_dr',  convert_to_pdf=False)

        # generate_content("2024-08-15-coracao.md", empresa='audaz', title='Sobre o coração', directory='data\\teste_exemplo\\input',  convert_to_pdf=True, create_cover=True, create_toc=True)
        # generate_content("exemplo.md", empresa='audaz', title='Exemplo', directory='data\\teste_exemplo',  convert_to_pdf=False, create_cover=False, create_toc=False)
        # generate_content("proposta_assembleia.md", directory='data\\ten_meetings_modelo_proposta',  title='Software como Serviço (SaaS) para assembléia digital' , 
                        #  convert_to_pdf=False, create_cover=True, create_toc=False)
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='Gera documentos.')
    parser.add_argument("--file", "-f", required=False)
    parser.add_argument("--empresa", "-e", required=False)
    parser.add_argument("--convert_to_pd", required=False)
    parser.add_argument("--generate_cover", required=False)
    parser.add_argument("--create_toc", required=False)
    parser.add_argument("-a", required=False)
    parser.add_argument("-b", required=False)
    # parser.add_argument("filename", ..., required=True)
    args = parser.parse_args()    
    print(f"args:{args}")
    print(f"args.file:{args.file}")
    main(args)
    exit(0)