from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor; 


# Create a new Word document
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

#title_style = document.styles['Normal']
#title_style.font.name = "Times New Roman"
#title_style.element.xml

heading0 = document.add_heading('AUDAZ - Infraestrutura em Nuvem ', 0 )
title_style = heading0.style
title_style.font.name = "Arial"
title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
title_style.font.bold = True
title_style.font.size = Pt(14)
title_style.element.xml
rFonts = title_style.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")

document.add_paragraph('p1.')
heading1 = document.add_heading('Arquitetura ', 1 )
title_style = heading1.style
title_style.font.name = "Arial"
title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
title_style.font.bold = True
title_style.font.size = Pt(12)
title_style.element.xml
rFonts = title_style.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")



document.add_paragraph('p2.')
heading2 = document.add_heading('Rede ', 2 )
title_style = heading2.style
title_style.font.name = "Arial"
title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
title_style.font.bold = True
title_style.font.size = Pt(11)
title_style.element.xml
rFonts = title_style.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")




document.add_paragraph('p3.')
document.add_page_break()

# Add a title to the document
document.add_heading('As-built Document of XCorp\'s Cloud Infrastructure', 0)

# Add a paragraph to the document
document.add_paragraph('This document provides an overview of XCorp\'s cloud infrastructure, which includes Kubernetes, KEDA, NGINX Ingress Controller, AWS CloudFront, AWS S3, AWS ECR, AWS Router53, Oracle Cloud Instance, RabbitMQ, Redis for caching, MySQL Heatwave, Postgres Database, ElasticSearch for logging and APM, and PRTG.')

# Add a new section to the document
document.add_heading('1. Architecture', level=1)

# Add a paragraph to the section
document.add_paragraph('XCorp\'s cloud infrastructure is designed to be highly available, scalable, and secure. The infrastructure is built on top of AWS and Oracle Cloud, with Kubernetes serving as the container orchestration platform.')

# Add a sub-section to the section
document.add_heading('1.1 Kubernetes', level=2)
document.add_paragraph('Kubernetes is used to manage and orchestrate containers across multiple nodes in the cluster. The cluster consists of multiple worker nodes, each running multiple pods, which are the smallest deployable units of computing in Kubernetes.')

# Add another sub-section to the section
document.add_heading('1.2 KEDA', level=2)
document.add_paragraph('KEDA (Kubernetes Event-driven Autoscaling) is used to automatically scale pods based on the number of events or messages received by the pod. This allows for fine-grained control over the scaling of applications and ensures that resources are used efficiently.')

# Add another sub-section to the section
document.add_heading('1.3 NGINX Ingress Controller', level=2)
document.add_paragraph('NGINX Ingress Controller is used to route traffic to the appropriate pods in the cluster. It acts as a reverse proxy and load balancer, ensuring that traffic is distributed evenly across the cluster.')

# Add another sub-section to the section
document.add_heading('1.4 AWS CloudFront', level=2)
document.add_paragraph('AWS CloudFront is used as a content delivery network (CDN) to cache static content and serve it to users from the nearest edge location. This helps to improve the performance of the application and reduce the load on the backend servers.')

# Add another sub-section to the section
document.add_heading('1.5 AWS S3', level=2)
document.add_paragraph('AWS S3 is used as a storage service for static content, such as images and videos. It provides high availability and durability, and can be used to store large amounts of data at a low cost.')

# Add another sub-section to the section
document.add_heading('1.6 AWS ECR', level=2)
document.add_paragraph('AWS ECR (Elastic Container Registry) is used as a container registry to store and manage container images. It provides secure and scalable storage for container images, and can be used to deploy containers to the cluster.')

# Add another sub-section to the section
document.add_heading('1.7 AWS Router53', level=2)
document.add_paragraph('AWS Router53 is used as a DNS service to route traffic to the appropriate IP address. It provides high availability and scalability, and can be used to manage DNS records for the application.')

# Add another sub-section to the section
document.add_heading('1.8 Oracle Cloud Instance', level=2)
document.add_paragraph('Oracle Cloud Instance is used as a compute service to run the application. It provides high performance and scalability, and can be used to run workloads that require a high degree of isolation and control.')

# Add another sub-section to the section
document.add_heading('1.9 RabbitMQ', level=2)
document.add_paragraph('RabbitMQ is used as a message broker to enable communication between different components of the application. It provides reliable and scalable messaging, and can be used to implement a variety of messaging patterns.')

# Add another sub-section to the section
#document.add_heading('1.10 Redis for caching', level=2

document.save('doc2.docx')