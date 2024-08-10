from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor; 

# Create a new Word document
document = Document()

# Set font and size for paragraphs
document.styles['Normal'].font.name = 'Arial'
document.styles['Normal'].font.size = Pt(10)
document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

# Set font and size for titles and subtitles
document.styles['Heading 1'].font.name = 'Arial'
document.styles['Heading 1'].font.size = Pt(14)
document.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
document.styles['Heading 2'].font.name = 'Arial'
document.styles['Heading 2'].font.size = Pt(12)
document.styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)

# Add the title
document.add_heading('As-built Document of XCorp\'s Cloud Infrastructure', 0)

# Add the subtitle
document.add_paragraph('This document provides an overview of XCorp\'s cloud infrastructure, which includes Kubernetes, KEDA, NGINX Ingress Controller, AWS CloudFront, AWS S3, AWS ECR, AWS Router53, Oracle Cloud Instance, RabbitMQ, Redis for caching, MySQL Heatwave, Postgres Database, ElasticSearch for logging and APM, and PRTG.', style='Heading 2')

# Add the section title
document.add_paragraph('1. Architecture', style='Heading 1')

# Add the sub-section titles
document.add_paragraph('1.1 Kubernetes', style='Heading 2')
document.add_paragraph('1.2 KEDA', style='Heading 2')
document.add_paragraph('1.3 NGINX Ingress Controller', style='Heading 2')
document.add_paragraph('1.4 AWS CloudFront', style='Heading 2')
document.add_paragraph('1.5 AWS S3', style='Heading 2')
document.add_paragraph('1.6 AWS ECR', style='Heading 2')
document.add_paragraph('1.7 AWS Router53', style='Heading 2')
document.add_paragraph('1.8 Oracle Cloud Instance', style='Heading 2')
document.add_paragraph('1.9 RabbitMQ', style='Heading 2')
document.add_paragraph('1.10 Redis for caching', style='Heading 2')
document.add_paragraph('1.11 MySQL Heatwave', style='Heading 2')
document.add_paragraph('1.12 Postgres Database', style='Heading 2')
document.add_paragraph('1.13 ElasticSearch for logging and APM', style='Heading 2')
document.add_paragraph('1.14 PRTG', style='Heading 2')

# Save the document
document.save('XCorp_Cloud_Infrastructure.docx')